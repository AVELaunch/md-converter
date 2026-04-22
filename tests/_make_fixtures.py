#!/usr/bin/env python3
"""
Generate deterministic test fixtures for converter round-trip tests.
Run once:  python3 tests/_make_fixtures.py
"""

from pathlib import Path

FIXTURES = Path(__file__).resolve().parent / "fixtures"
FIXTURES.mkdir(exist_ok=True)


def make_txt():
    (FIXTURES / "sample.txt").write_text("Hello world\nThis is a test.", encoding="utf-8")


def make_html():
    (FIXTURES / "sample.html").write_text(
        "<html><head><title>Test</title></head>"
        "<body><h1>H</h1><p>Para.</p></body></html>",
        encoding="utf-8",
    )


def make_docx():
    from docx import Document

    doc = Document()
    doc.add_heading("Sample Heading", level=1)
    doc.add_paragraph("This is a sample paragraph for testing.")
    doc.save(str(FIXTURES / "sample.docx"))


def make_pdf():
    import fitz

    doc = fitz.open()
    for text in ("Page one text content here.", "Page two text content here."):
        page = doc.new_page()
        page.insert_text((72, 72), text, fontsize=12)
    doc.save(str(FIXTURES / "sample.pdf"))
    doc.close()


def make_scanned_pdf():
    """Create a PDF whose pages are images (no selectable text).

    PyMuPDF text extraction returns empty for these pages,
    triggering the OCR fallback path.
    """
    import fitz

    # Create a temporary PDF with text, render pages to PNG images,
    # then embed those PNGs in a new PDF with NO text layer.
    # PyMuPDF get_text() returns empty → triggers OCR fallback.
    tmp = fitz.open()
    for text in ("Scanned page one.", "Scanned page two."):
        page = tmp.new_page()
        page.insert_text((72, 72), text, fontsize=14)

    doc = fitz.open()
    for i in range(len(tmp)):
        pix = tmp[i].get_pixmap(dpi=150)
        png_bytes = pix.tobytes("png")
        page = doc.new_page(width=pix.width * 72 / 150, height=pix.height * 72 / 150)
        page.insert_image(page.rect, stream=png_bytes)
    tmp.close()

    doc.save(str(FIXTURES / "sample_scanned.pdf"))
    doc.close()


def make_rtf():
    (FIXTURES / "sample.rtf").write_text(
        r"{\rtf1\ansi Hello.}", encoding="utf-8"
    )


if __name__ == "__main__":
    make_txt()
    print("  created sample.txt")
    make_html()
    print("  created sample.html")
    make_docx()
    print("  created sample.docx")
    make_pdf()
    print("  created sample.pdf")
    make_scanned_pdf()
    print("  created sample_scanned.pdf")
    make_rtf()
    print("  created sample.rtf")
    print("Done.")
