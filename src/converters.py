#!/usr/bin/env python3
"""
Universal Markdown Converters
Five format converters with shared utilities for consistent output.
"""

import ipaddress
import json as _json
import logging
import os
import re
import shutil
import socket
import ssl
import sys
import time
from datetime import date
from pathlib import Path
from typing import NamedTuple

logger = logging.getLogger("md_converter.converters")

import certifi

# Fix SSL certificate discovery on macOS + Python 3.14 + OpenSSL 3.6:
# certifi's bundled CA file may be rejected by newer OpenSSL builds.
# Prefer the system/Homebrew CA file that OpenSSL ships with; fall back
# to certifi only when no system file is found.
def _find_ca_file() -> str:
    """Return the best available CA certificate bundle path."""
    # 1. Already set by user/environment
    env_ca = os.environ.get("SSL_CERT_FILE")
    if env_ca and os.path.isfile(env_ca):
        return env_ca
    # 2. System OpenSSL default (works with Homebrew OpenSSL 3.x)
    _paths = ssl.get_default_verify_paths()
    for candidate in (_paths.cafile, _paths.openssl_cafile):
        if candidate and os.path.isfile(candidate):
            return candidate
    # 3. Fallback to certifi
    return certifi.where()

_CA_FILE = _find_ca_file()

import fitz  # PyMuPDF
import requests

# Maximum response size for URL fetches (50 MB)
_MAX_RESPONSE_BYTES = 50 * 1024 * 1024


class FetchResult(NamedTuple):
    content: bytes
    encoding: str


def _is_private_ip(host: str) -> bool:
    """Check if any resolved IP for *host* is private/reserved."""
    try:
        infos = socket.getaddrinfo(host, None, proto=socket.IPPROTO_TCP)
    except socket.gaierror:
        raise ValueError(f"Cannot resolve hostname: {host}")
    for info in infos:
        addr = ipaddress.ip_address(info[4][0])
        if addr.is_private or addr.is_loopback or addr.is_link_local or addr.is_reserved or addr.is_multicast:
            return True
    return False


def _safe_get(url: str) -> FetchResult:
    """Fetch *url* with SSRF guards, redirect validation, and size cap."""
    from urllib.parse import urlparse

    parsed = urlparse(url)
    if parsed.scheme not in ("http", "https"):
        raise ValueError(f"Unsupported URL scheme: {parsed.scheme!r}")
    if not parsed.hostname:
        raise ValueError("URL has no hostname")
    if _is_private_ip(parsed.hostname):
        raise ValueError(f"Request to private/reserved address blocked: {parsed.hostname}")

    # Set SSL_CERT_FILE only when we actually make a request
    os.environ["SSL_CERT_FILE"] = _CA_FILE

    resp = requests.get(
        url,
        stream=True,
        timeout=30,
        headers={"User-Agent": "MDConverter/1.0"},
        verify=_CA_FILE,
        allow_redirects=True,
    )
    resp.raise_for_status()

    # Validate every redirect hop and the final landed URL
    for hop in resp.history:
        hop_host = urlparse(hop.url).hostname
        if hop_host and _is_private_ip(hop_host):
            raise ValueError("Redirect to private address blocked")
    final_host = urlparse(str(resp.url)).hostname
    if final_host and _is_private_ip(final_host):
        raise ValueError("Redirect to private address blocked")

    # Read with size cap
    chunks = []
    total = 0
    for chunk in resp.iter_content(chunk_size=65536):
        total += len(chunk)
        if total > _MAX_RESPONSE_BYTES:
            resp.close()
            raise ValueError(f"Response exceeds {_MAX_RESPONSE_BYTES // (1024 * 1024)} MB limit")
        chunks.append(chunk)

    return FetchResult(
        content=b"".join(chunks),
        encoding=resp.encoding or "utf-8",
    )
from bs4 import BeautifulSoup
from docx import Document
from docx.table import Table as DocxTable
from markdownify import markdownify as html_to_md
from striprtf.striprtf import rtf_to_text


# ---------------------------------------------------------------------------
# Shared types
# ---------------------------------------------------------------------------

class ConvertResult(NamedTuple):
    success: bool
    output_path: str
    word_count: int
    message: str


# ---------------------------------------------------------------------------
# Shared utilities (extracted from pdf_to_md / docx_to_md)
# ---------------------------------------------------------------------------

def safe_filename(name: str) -> str:
    """Sanitize a document title into a lowercase-dash filename."""
    safe = re.sub(r'[^\w\s\-]', '', name).strip()
    return re.sub(r'\s+', '-', safe).lower()


def normalize_blanks(text: str) -> str:
    """Collapse 3+ consecutive newlines into 2."""
    return re.sub(r'\n{3,}', '\n\n', text)


def build_header(title: str, source: str, word_count: int, **extras) -> str:
    """Build the blockquote metadata header used by all converters."""
    lines = [f"# {title}", ""]
    lines.append(f"> **Source**: `{source}`  ")
    for key, val in extras.items():
        lines.append(f"> **{key}**: {val}  ")
    lines.append(f"> **Word Count**: {word_count:,}  ")
    lines.extend(["", "---", ""])
    return "\n".join(lines)


def _yaml_safe(value: str) -> str:
    """Return *value* quoted safely for use as a YAML scalar.

    If the value contains characters that could break YAML parsing
    (quotes, backslashes, newlines, or leading special chars), wrap it
    with json.dumps which produces a valid YAML double-quoted string.
    """
    if any(c in value for c in ('"', '\\', '\n', '\r')) or (value and value[0] in '&*?|>{[!%@`'):
        return _json.dumps(value)
    return f'"{value}"'


def vault_frontmatter(title: str, source_type: str, source_file: str) -> str:
    """Generate Obsidian-compatible YAML frontmatter."""
    return (
        f"---\n"
        f"title: {_yaml_safe(title)}\n"
        f"tags: [converted, {source_type}]\n"
        f"created: {date.today().isoformat()}\n"
        f"source: {_yaml_safe(source_file)}\n"
        f"---\n\n"
    )


def write_output(
    body: str,
    title: str,
    source_file: str,
    word_count: int,
    output_dir: Path,
    source_type: str,
    vault_dir: Path | None = None,
    header_extras: dict | None = None,
) -> ConvertResult:
    """Write markdown to output_dir and optionally copy to vault."""
    output_dir.mkdir(parents=True, exist_ok=True)
    md_name = f"{safe_filename(title)}.md"
    md_path = output_dir / md_name

    header = build_header(title, source_file, word_count, **(header_extras or {}))
    content = header + body + "\n"
    md_path.write_text(content, encoding="utf-8")

    # Vault delivery
    if vault_dir:
        vault_type_dir = vault_dir / source_type
        vault_type_dir.mkdir(parents=True, exist_ok=True)
        vault_path = vault_type_dir / md_name
        vault_content = vault_frontmatter(title, source_type, source_file) + content
        vault_path.write_text(vault_content, encoding="utf-8")

    return ConvertResult(True, str(md_path), word_count, f"OK -> {md_name}")


# ---------------------------------------------------------------------------
# Marker converter cache
# ---------------------------------------------------------------------------

_MARKER_CONVERTER = None
_MARKER_LOADED = False


def _get_marker_converter():
    """Return a cached PdfConverter, loading models on first call."""
    global _MARKER_CONVERTER, _MARKER_LOADED
    if _MARKER_CONVERTER is None:
        from marker.converters.pdf import PdfConverter
        from marker.models import create_model_dict
        if not _MARKER_LOADED:
            logger.info("Loading Marker models (first run downloads ~1 GB)...")
        else:
            logger.debug("Reusing cached Marker converter")
        _MARKER_CONVERTER = PdfConverter(artifact_dict=create_model_dict())
        _MARKER_LOADED = True
    return _MARKER_CONVERTER


# ---------------------------------------------------------------------------
# Config path resolution
# ---------------------------------------------------------------------------

def _config_path() -> Path:
    """Return the config.json path, respecting frozen builds."""
    if getattr(sys, "frozen", False):
        base = Path.home() / "Library" / "Application Support" / "MD Converter"
        base.mkdir(parents=True, exist_ok=True)
        return base / "config.json"
    return Path(__file__).resolve().parent.parent / "config.json"


# ---------------------------------------------------------------------------
# OCR engine selection
# ---------------------------------------------------------------------------

# Priority: MD_CONVERTER_OCR env var > config.json "ocr_engine" > "auto"
OCR_ENGINE: str | None = os.environ.get("MD_CONVERTER_OCR", "").lower() or None


def _get_ocr_engine() -> str:
    """Return the effective OCR engine name."""
    if OCR_ENGINE is not None:
        return OCR_ENGINE
    config = _config_path()
    if config.exists():
        try:
            cfg = _json.loads(config.read_text(encoding="utf-8"))
            engine = cfg.get("ocr_engine", "auto")
            if isinstance(engine, str):
                return engine.lower()
        except (ValueError, OSError):
            pass
    return "auto"


# ---------------------------------------------------------------------------
# Format routing
# ---------------------------------------------------------------------------

SUPPORTED = {'.pdf', '.docx', '.html', '.htm', '.txt', '.rtf'}

SUBFOLDER = {
    '.pdf': 'pdf', '.docx': 'docx',
    '.html': 'html', '.htm': 'html',
    '.txt': 'txt', '.rtf': 'rtf',
}


def route(path: str, base_output: Path, vault_dir: Path | None = None) -> ConvertResult:
    """Detect format and call the right converter."""
    # URL detection
    if path.startswith("http://") or path.startswith("https://"):
        out = base_output / "html"
        return convert_html(path, out, vault_dir)

    p = Path(path)
    ext = p.suffix.lower()
    if ext not in SUPPORTED:
        return ConvertResult(False, "", 0, f"Unsupported format: {ext}")

    out = base_output / SUBFOLDER[ext]
    converters = {
        '.pdf': convert_pdf,
        '.docx': convert_docx,
        '.html': convert_html,
        '.htm': convert_html,
        '.txt': convert_txt,
        '.rtf': convert_rtf,
    }
    return converters[ext](path, out, vault_dir)


# ---------------------------------------------------------------------------
# 1. PDF converter (with OCR auto-fallback)
# ---------------------------------------------------------------------------

def convert_pdf(path: str, output_dir: Path, vault_dir: Path | None = None) -> ConvertResult:
    """Convert a PDF to Markdown. Falls back to OCR if no selectable text."""
    p = Path(path)
    doc = fitz.open(path)
    pages = []
    total_words = 0
    page_count = len(doc)

    for i in range(page_count):
        text = doc[i].get_text("text")
        if text.strip():
            pages.append((i + 1, text))
            total_words += len(text.split())
    doc.close()

    # Auto-fallback to OCR
    if total_words == 0:
        engine = _get_ocr_engine()
        logger.info("No selectable text in %s — OCR engine: %s", Path(path).name, engine)
        if engine == "tesseract":
            return _convert_pdf_tesseract(path, output_dir, vault_dir)
        elif engine == "marker":
            return _convert_pdf_marker(path, output_dir, vault_dir)
        else:  # "auto" — try Marker first, fall back to Tesseract
            try:
                result = _convert_pdf_marker(path, output_dir, vault_dir)
                if result.success:
                    return result
                logger.warning("Marker returned empty for %s, falling back to Tesseract", Path(path).name)
            except Exception as exc:
                logger.warning("Marker failed for %s (%s), falling back to Tesseract", Path(path).name, exc)
            return _convert_pdf_tesseract(path, output_dir, vault_dir)

    body_lines = []
    for num, text in pages:
        body_lines.append(f"## Page {num}\n")
        body_lines.append(normalize_blanks(text.strip()))
        body_lines.append("")

    body = "\n".join(body_lines)
    return write_output(
        body, p.stem, p.name, total_words, output_dir, "pdf", vault_dir,
        header_extras={"Pages": str(page_count)},
    )


def _convert_pdf_marker(path: str, output_dir: Path, vault_dir: Path | None = None) -> ConvertResult:
    """OCR via Marker (local deep-learning models)."""
    try:
        from marker.output import text_from_rendered
    except ImportError:
        return ConvertResult(False, "", 0, "ERROR: marker-pdf not installed")

    start = time.time()
    converter = _get_marker_converter()
    rendered = converter(path)
    body, _metadata, _images = text_from_rendered(rendered)
    elapsed = time.time() - start
    word_count = len(body.split())
    logger.info("Marker finished in %.1fs — %d words extracted", elapsed, word_count)

    if word_count == 0:
        return ConvertResult(False, "", 0, "Marker returned empty")

    p = Path(path)
    with fitz.open(path) as _doc:
        page_count = len(_doc)
    return write_output(
        body, p.stem, p.name, word_count, output_dir, "pdf", vault_dir,
        header_extras={
            "Pages": str(page_count),
            "Extracted via": "Marker (local)",
            "Processing time": f"{elapsed:.1f}s",
        },
    )


def _convert_pdf_tesseract(path: str, output_dir: Path, vault_dir: Path | None = None) -> ConvertResult:
    """OCR fallback for scanned PDFs using Tesseract."""
    try:
        import pytesseract
        from PIL import Image
    except ImportError:
        return ConvertResult(False, "", 0, "ERROR: pytesseract/Pillow not installed for OCR")

    p = Path(path)
    doc = fitz.open(path)
    page_count = len(doc)
    pages = []
    total_words = 0
    zoom = 300 / 72
    matrix = fitz.Matrix(zoom, zoom)
    start = time.time()

    for i in range(page_count):
        pix = doc[i].get_pixmap(matrix=matrix)
        img = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)
        text = pytesseract.image_to_string(img)
        if text.strip():
            pages.append((i + 1, text))
            total_words += len(text.split())
    doc.close()
    elapsed = time.time() - start

    if total_words == 0:
        return ConvertResult(False, "", 0, "ERROR: No text extracted even with OCR")

    body_lines = []
    for num, text in pages:
        body_lines.append(f"## Page {num}\n")
        body_lines.append(normalize_blanks(text.strip()))
        body_lines.append("")

    body = "\n".join(body_lines)
    return write_output(
        body, p.stem, p.name, total_words, output_dir, "pdf", vault_dir,
        header_extras={
            "Pages": str(page_count),
            "Extracted via": "Tesseract OCR (local)",
            "Processing time": f"{elapsed:.1f}s",
        },
    )


# ---------------------------------------------------------------------------
# 2. DOCX converter (ported from docx_to_md.py)
# ---------------------------------------------------------------------------

def _run_to_md(run) -> str:
    text = run.text
    if not text:
        return ""
    if run.bold and run.italic:
        return f"***{text}***"
    if run.bold:
        return f"**{text}**"
    if run.italic:
        return f"*{text}*"
    return text


def _para_to_md(para) -> str:
    style = para.style.name.lower()
    parts = [_run_to_md(r) for r in para.runs]
    text = "".join(parts).strip() or para.text.strip()
    if not text:
        return ""

    if style.startswith("heading"):
        try:
            level = min(int(style.split()[-1]), 6)
        except (ValueError, IndexError):
            level = 1
        return f"{'#' * level} {text}"
    if style == "title":
        return f"# {text}"
    if style == "subtitle":
        return f"## {text}"
    if style.startswith("list bullet"):
        depth = style.count("2") + style.count("3")
        return f"{'  ' * depth}- {text}"
    if style.startswith("list number"):
        depth = style.count("2") + style.count("3")
        return f"{'  ' * depth}1. {text}"
    if "quote" in style:
        return f"> {text}"
    return text


def _table_to_md(table: DocxTable) -> str:
    rows = []
    for row in table.rows:
        cells = [c.text.strip().replace("\n", " ").replace("|", "\\|") for c in row.cells]
        rows.append(cells)
    if not rows:
        return ""
    cols = max(len(r) for r in rows)
    for r in rows:
        while len(r) < cols:
            r.append("")
    lines = ["| " + " | ".join(rows[0]) + " |"]
    lines.append("| " + " | ".join(["---"] * cols) + " |")
    for row in rows[1:]:
        lines.append("| " + " | ".join(row) + " |")
    return "\n".join(lines)


def convert_docx(path: str, output_dir: Path, vault_dir: Path | None = None) -> ConvertResult:
    """Convert a DOCX file preserving headings, formatting, lists, and tables."""
    p = Path(path)
    doc = Document(path)
    blocks = []
    word_count = 0

    for child in doc.element.body:
        tag = child.tag.split("}")[-1]
        if tag == "p":
            for para in doc.paragraphs:
                if para._element is child:
                    line = _para_to_md(para)
                    blocks.append(line if line else "")
                    if line:
                        word_count += len(line.split())
                    break
        elif tag == "tbl":
            for table in doc.tables:
                if table._tbl is child:
                    md = _table_to_md(table)
                    if md:
                        blocks.extend(["", md, ""])
                        word_count += len(md.split())
                    break

    body = normalize_blanks("\n".join(blocks)).strip()
    if word_count == 0:
        return ConvertResult(False, "", 0, "SKIPPED (empty)")

    return write_output(body, p.stem, p.name, word_count, output_dir, "docx", vault_dir)


# ---------------------------------------------------------------------------
# 3. HTML / URL converter
# ---------------------------------------------------------------------------

def convert_html(path_or_url: str, output_dir: Path, vault_dir: Path | None = None) -> ConvertResult:
    """Convert HTML file or URL to Markdown."""
    is_url = path_or_url.startswith("http://") or path_or_url.startswith("https://")

    if is_url:
        result = _safe_get(path_or_url)
        html = result.content.decode(result.encoding, errors="replace")
        soup = BeautifulSoup(html, "html.parser")
        title = soup.title.string.strip() if soup.title and soup.title.string else path_or_url
        source = path_or_url
    else:
        p = Path(path_or_url)
        html = p.read_text(encoding="utf-8", errors="replace")
        soup = BeautifulSoup(html, "html.parser")
        title = soup.title.string.strip() if soup.title and soup.title.string else p.stem
        source = p.name

    # Remove script/style tags before conversion
    for tag in soup(["script", "style", "nav", "footer", "header"]):
        tag.decompose()

    body = html_to_md(str(soup), heading_style="ATX", strip=["img"])
    body = normalize_blanks(body).strip()
    word_count = len(body.split())

    if word_count == 0:
        return ConvertResult(False, "", 0, "SKIPPED (empty page)")

    extras = {"URL": f"`{path_or_url}`"} if is_url else {}
    return write_output(body, title, source, word_count, output_dir, "html", vault_dir, extras)


# ---------------------------------------------------------------------------
# 4. TXT converter
# ---------------------------------------------------------------------------

def convert_txt(path: str, output_dir: Path, vault_dir: Path | None = None) -> ConvertResult:
    """Wrap a plain text file in a Markdown metadata header."""
    p = Path(path)
    text = p.read_text(encoding="utf-8", errors="replace").strip()
    word_count = len(text.split())

    if word_count == 0:
        return ConvertResult(False, "", 0, "SKIPPED (empty)")

    return write_output(text, p.stem, p.name, word_count, output_dir, "txt", vault_dir)


# ---------------------------------------------------------------------------
# 5. RTF converter
# ---------------------------------------------------------------------------

def convert_rtf(path: str, output_dir: Path, vault_dir: Path | None = None) -> ConvertResult:
    """Convert RTF to Markdown by stripping RTF formatting."""
    p = Path(path)
    try:
        # Read as raw bytes first — RTF files are almost never UTF-8.
        raw_bytes = p.read_bytes()

        # Decode with latin-1, which maps every byte 0x00-0xFF one-to-one to
        # Unicode code-points. This preserves the original bytes so that
        # striprtf can interpret RTF encoding directives (\\ansicpg, etc.)
        # and produce correct Unicode output.
        raw = raw_bytes.decode("latin-1")

        text = rtf_to_text(raw, errors="ignore").strip()
    except Exception as e:
        return ConvertResult(False, "", 0, f"ERROR reading RTF: {e}")

    word_count = len(text.split())

    if word_count == 0:
        return ConvertResult(False, "", 0, "SKIPPED (empty)")

    return write_output(text, p.stem, p.name, word_count, output_dir, "rtf", vault_dir)


# ---------------------------------------------------------------------------
# 6. Raw pasted text converter
# ---------------------------------------------------------------------------

def convert_raw_text(text: str, output_dir: Path, vault_dir: Path | None = None) -> ConvertResult:
    """Save raw pasted text as a Markdown file."""
    text = text.strip()
    word_count = len(text.split())
    if word_count == 0:
        return ConvertResult(False, "", 0, "SKIPPED (empty)")

    # Generate title from first line (truncated to 60 chars)
    first_line = text.split('\n')[0].strip()
    title = first_line[:60] if first_line else "Pasted Text"
    # Clean title for display
    title = re.sub(r'[#*>\-=]', '', title).strip() or "Pasted Text"

    return write_output(text, title, "pasted-text", word_count, output_dir, "txt", vault_dir)


# ---------------------------------------------------------------------------
# 7. Pasted input router (text or URL)
# ---------------------------------------------------------------------------

def convert_pasted(text: str, base_output: Path, vault_dir: Path | None = None) -> ConvertResult:
    """Route pasted text - if it looks like a URL, fetch it; otherwise save as text."""
    text = text.strip()
    if text.startswith("http://") or text.startswith("https://"):
        out = base_output / "html"
        return convert_html(text, out, vault_dir)
    return convert_raw_text(text, base_output / "txt", vault_dir)
