#!/usr/bin/env python3
"""
DOCX to Markdown Converter
Converts all .docx files in a directory to Markdown with full formatting
preservation — headings, bold, italic, lists, tables, and word count summaries.
Output goes to converted/docx/ to separate from PDF conversions.
"""

import re
import sys
from pathlib import Path

from docx import Document
from docx.table import Table


def run_to_markdown(run) -> str:
    """Convert a single text run to Markdown with inline formatting."""
    text = run.text
    if not text:
        return ""
    if run.bold and run.italic:
        return f"***{text}***"
    if run.bold:
        return f"**{text}**"
    if run.italic:
        return f"*{text}*"
    return text


def paragraph_to_markdown(para) -> str:
    """Convert a paragraph to a Markdown line with heading/list detection."""
    style = para.style.name.lower()

    # Build inline-formatted text from runs
    parts = [run_to_markdown(r) for r in para.runs]
    text = "".join(parts).strip()

    # Fall back to plain text if runs produced nothing
    if not text:
        text = para.text.strip()
    if not text:
        return ""

    # Headings
    if style.startswith("heading"):
        try:
            level = int(style.split()[-1])
        except (ValueError, IndexError):
            level = 1
        level = min(level, 6)
        return f"{'#' * level} {text}"

    if style == "title":
        return f"# {text}"
    if style == "subtitle":
        return f"## {text}"

    # Bullet lists
    if style.startswith("list bullet"):
        depth = style.count("2") + style.count("3")
        return f"{'  ' * depth}- {text}"

    # Numbered lists
    if style.startswith("list number"):
        depth = style.count("2") + style.count("3")
        return f"{'  ' * depth}1. {text}"

    # Block quotes
    if "quote" in style:
        return f"> {text}"

    return text


def table_to_markdown(table: Table) -> str:
    """Convert a DOCX table to a Markdown table."""
    rows = []
    for row in table.rows:
        cells = [cell.text.strip().replace("\n", " ") for cell in row.cells]
        rows.append(cells)

    if not rows:
        return ""

    col_count = max(len(r) for r in rows)
    for r in rows:
        while len(r) < col_count:
            r.append("")

    lines = []
    lines.append("| " + " | ".join(rows[0]) + " |")
    lines.append("| " + " | ".join(["---"] * col_count) + " |")
    for row in rows[1:]:
        lines.append("| " + " | ".join(row) + " |")

    return "\n".join(lines)


def extract_docx(docx_path: str) -> tuple[str, int]:
    """
    Extract a DOCX file into Markdown text.
    Returns (markdown_body, word_count).
    """
    doc = Document(docx_path)
    blocks = []
    word_count = 0

    body = doc.element.body
    for child in body:
        tag = child.tag.split("}")[-1]

        if tag == "p":
            for para in doc.paragraphs:
                if para._element is child:
                    line = paragraph_to_markdown(para)
                    if line:
                        blocks.append(line)
                        word_count += len(line.split())
                    else:
                        blocks.append("")
                    break

        elif tag == "tbl":
            for table in doc.tables:
                if table._tbl is child:
                    md_table = table_to_markdown(table)
                    if md_table:
                        blocks.append("")
                        blocks.append(md_table)
                        blocks.append("")
                        word_count += len(md_table.split())
                    break

    content = "\n".join(blocks)
    content = re.sub(r'\n{3,}', '\n\n', content)
    return content.strip(), word_count


def build_markdown(body: str, title: str, word_count: int, source_file: str) -> str:
    """Wrap extracted body in a metadata header."""
    lines = [
        f"# {title}",
        "",
        f"> **Source**: `{source_file}`  ",
        f"> **Word Count**: {word_count:,}  ",
        "",
        "---",
        "",
        body,
        "",
    ]
    return "\n".join(lines)


def convert_docx_files(source_dir: str, output_dir: str):
    """Find and convert all .docx files in source_dir, saving to output_dir."""
    source_path = Path(source_dir)
    output_path = Path(output_dir)
    output_path.mkdir(parents=True, exist_ok=True)

    docx_files = sorted(source_path.rglob("*.docx"))
    docx_files = [f for f in docx_files if not f.name.startswith("~$")]

    if not docx_files:
        print("No .docx files found.")
        return

    print(f"Found {len(docx_files)} DOCX file(s)\n")
    print(f"{'File':<60} {'Words':>10} {'Status'}")
    print("-" * 90)

    results = []

    for docx_file in docx_files:
        name = docx_file.stem
        relative = docx_file.relative_to(source_path)
        display = str(relative)[:57] + "..." if len(str(relative)) > 60 else str(relative)

        try:
            body, word_count = extract_docx(str(docx_file))

            if word_count == 0:
                status = "SKIPPED (empty)"
                print(f"{display:<60} {word_count:>10,} {status}")
                results.append((str(relative), word_count, status))
                continue

            safe_name = re.sub(r'[^\w\s\-]', '', name).strip()
            safe_name = re.sub(r'\s+', '-', safe_name).lower()
            md_filename = f"{safe_name}.md"
            md_path = output_path / md_filename

            md_content = build_markdown(body, name, word_count, str(relative))
            md_path.write_text(md_content, encoding="utf-8")

            status = f"OK -> {md_filename}"
            print(f"{display:<60} {word_count:>10,} {status}")
            results.append((str(relative), word_count, status))

        except Exception as e:
            status = f"ERROR: {e}"
            print(f"{display:<60} {'?':>10} {status}")
            results.append((str(relative), 0, status))

    ok_count = len([r for r in results if "OK" in r[2]])
    total_words = sum(r[1] for r in results)

    print("\n" + "=" * 90)
    print(f"SUMMARY: {ok_count}/{len(docx_files)} files converted | {total_words:,} total words")
    print(f"Output directory: {output_path.resolve()}")


if __name__ == "__main__":
    script_dir = Path(__file__).parent
    source_dir = script_dir / "Docs"
    output_dir = script_dir / "converted" / "docx"

    if len(sys.argv) >= 2:
        source_dir = Path(sys.argv[1])
    if len(sys.argv) >= 3:
        output_dir = Path(sys.argv[2])

    convert_docx_files(str(source_dir), str(output_dir))
